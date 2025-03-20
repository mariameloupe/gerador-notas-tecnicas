import streamlit as st
import pandas as pd
import os
from docx import Document
from datetime import datetime
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Caminho fixo da logo
LOGO_PATH = "logo.png"  # Certifique-se de que a logo está salva neste caminho

# Função para carregar a planilha de pagamentos
def carregar_dados(uploaded_file, tipo_planilha):
    if uploaded_file is not None:
        xls = pd.ExcelFile(uploaded_file)
        
        if tipo_planilha == "FEM":
            df_inf_gerais = pd.read_excel(xls, "INF GERAIS", skiprows=5, dtype=str)
            df_inf_gerais.columns = df_inf_gerais.columns.str.strip()
            df_inf_gerais = df_inf_gerais[["ORDEM", "MUNICÍPIO", "PROJETO", "PROJETO DETALHADO", "TETO FEM", "STATUS PTM", "STATUS OBRA", "RESSALVA"]]
            
            df_resumo = pd.read_excel(xls, "RESUMO", skiprows=5, dtype=str)
            df_resumo.columns = df_resumo.columns.str.strip()
            df_resumo = df_resumo[["ORDEM", "MUNICÍPIO", "PROJETO", "STATUS PTM", "STATUS OBRA", "TETO FEM", "DATA ÚLTIMO PAGAMENTO", "REPASSE_VÁLIDO"]]
            
            df_inf_gerais["ORDEM"] = df_inf_gerais["ORDEM"].astype(str)
            df_resumo["ORDEM"] = df_resumo["ORDEM"].astype(str)
            df = pd.merge(df_inf_gerais, df_resumo, on="ORDEM", how="left")
            return df
        
        elif tipo_planilha == "EMENDAS":
            df_inf_gerais = pd.read_excel(xls, "INF GERAIS", skiprows=5, dtype=str)
            df_inf_gerais.columns = df_inf_gerais.columns.str.strip()
            df_inf_gerais = df_inf_gerais[["PROJETO DETALHADO"]]
            
            df_resumo = pd.read_excel(xls, "RESUMO", skiprows=5, dtype=str)
            df_resumo.columns = df_resumo.columns.str.strip()
            df_resumo = df_resumo[["MUNICÍPIO", "STATUS OBRA", "VALOR UTILIZADO DA EMENDA", "REPASSE_VÁLIDO", "DATA ÚLTIMO PAGAMENTO"]]
            
            df = pd.concat([df_inf_gerais, df_resumo], axis=1)
            return df
    
    return None

# Função para adicionar parágrafos formatados
def adicionar_paragrafo_formatado(doc, texto, fonte='Arial', tamanho=12, negrito=False, cor=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(texto)
    run.font.name = fonte
    run.font.size = Pt(tamanho)
    run.bold = negrito
    if cor:
        run.font.color.rgb = cor
    return paragraph

# Função para formatar valores como moeda
def formatar_moeda(valor):
    try:
        return f"R$ {float(valor):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except ValueError:
        return 'R$ 0,00'

# Função para aplicar fonte Arial tamanho 8 em todas as células da tabela
def formatar_tabela(tabela):
    for row in tabela.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)

# Função para definir a cor de fundo de uma célula
def definir_cor_fundo_celula(celula, cor_hex):
    """
    Define a cor de fundo de uma célula da tabela.
    cor_hex: Cor em formato hexadecimal (ex: #75B4FF).
    """
    cor_rgb = tuple(int(cor_hex.lstrip('#')[i:i+2], 16) for i in (0, 2, 4))
    tcPr = celula._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), cor_hex.lstrip('#'))
    tcPr.append(shd)

# Função para formatar o cabeçalho da tabela (negrito, cor branca e fundo colorido)
def formatar_cabecalho_tabela(celula, texto, cor_fundo='#75B4FF'):
    celula.text = texto
    for paragraph in celula.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.bold = True  # Texto em negrito
            run.font.color.rgb = RGBColor(255, 255, 255)  # Cor da fonte branca
    definir_cor_fundo_celula(celula, cor_fundo)  # Cor de fundo

# Função para formatar células em negrito
def formatar_celula_negrito(celula, texto):
    celula.text = texto
    for paragraph in celula.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.bold = True  # Texto em negrito

# Layout no Streamlit
st.set_page_config(page_title="Gerador de Notas Técnicas", layout="wide")

st.markdown("<h3 style='text-align: right;'>SEPLAG-SEDRC</h3>", unsafe_allow_html=True)

st.title("📄 Gerador de Notas Técnicas")
st.markdown("Este aplicativo gera automaticamente notas técnicas a partir das planilhas de controle do FEM e de Emendas Parlamentares.")

if os.path.exists(LOGO_PATH):
    st.sidebar.image(LOGO_PATH, use_container_width=True)

st.sidebar.header("Upload do Arquivo")
uploaded_files = st.sidebar.file_uploader("Faça upload dos arquivos Excel", type=["xlsm"], accept_multiple_files=True)

if uploaded_files:
    dfs_fem = []
    dfs_emendas = []
    for uploaded_file in uploaded_files:
        tipo_planilha = st.sidebar.selectbox(f"Tipo de planilha para {uploaded_file.name}", ["FEM", "EMENDAS"])
        df_temp = carregar_dados(uploaded_file, tipo_planilha)
        if df_temp is not None:
            if tipo_planilha == "FEM":
                dfs_fem.append(df_temp)
            else:
                dfs_emendas.append(df_temp)
    
    if dfs_fem:
        df_fem = pd.concat(dfs_fem, ignore_index=True)
        df_fem['ANO'] = df_fem['ORDEM'].astype(str).str[:4]  # Extrai o ano dos 4 primeiros dígitos da coluna ORDEM
    
    if dfs_emendas:
        df_emendas = pd.concat(dfs_emendas, ignore_index=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📊 Visualização da Planilha FEM")
        if dfs_fem:
            st.dataframe(df_fem, height=300)
        else:
            st.write("Nenhuma planilha FEM carregada.")
        
        st.subheader("📊 Visualização da Planilha de Emendas")
        if dfs_emendas:
            st.dataframe(df_emendas, height=300)
        else:
            st.write("Nenhuma planilha de Emendas carregada.")
    
    coluna_municipio_fem = 'MUNICÍPIO_x' if 'MUNICÍPIO_x' in df_fem.columns else 'MUNICÍPIO_y'
    coluna_municipio_emendas = 'MUNICÍPIO'
    
    municipios_disponiveis_fem = df_fem[coluna_municipio_fem].dropna().unique()
    municipios_disponiveis_emendas = df_emendas[coluna_municipio_emendas].dropna().unique()
    
    municipio_selecionado = st.selectbox("🌍 Selecione um Município", list(set(municipios_disponiveis_fem).union(set(municipios_disponiveis_emendas))))
    
    df_fem_filtrado = df_fem[df_fem[coluna_municipio_fem] == municipio_selecionado]
    df_emendas_filtrado = df_emendas[df_emendas[coluna_municipio_emendas] == municipio_selecionado]
    
    anos_disponiveis = df_fem_filtrado['ANO'].unique()
    data_atualizacao = st.text_input("📅 Atualizado em (Digite a data e hora)", value=datetime.today().strftime('%d/%m/%Y %H:%M'))
    
    if not df_fem_filtrado.empty and st.button("📝 Gerar Nota Técnica"):
        caminho_nota = "nota_tecnica.docx"
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(12)
        
        if os.path.exists(LOGO_PATH):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(LOGO_PATH, width=Inches(1.5))
            paragraph.alignment = 4  # Alinha à direita
        
        title = doc.add_paragraph()
        title_run = title.add_run('NOTA TÉCNICA')
        title_run.bold = True
        title.alignment = 1  # Centraliza o texto
        municipio_paragraph = doc.add_paragraph()
        municipio_run = municipio_paragraph.add_run(f"Município de {municipio_selecionado}")
        municipio_run.bold = True
        municipio_run.font.color.rgb = None
        doc.add_paragraph(f"Atualizado em: {data_atualizacao}")
        
        for ano in sorted(anos_disponiveis):
            adicionar_paragrafo_formatado(doc, f"Ano do Registro: {ano}", fonte='Arial', tamanho=12, negrito=False, cor=None)
            df_ano = df_fem_filtrado[df_fem_filtrado['ANO'] == ano]
            
            # Define os novos cabeçalhos
            colunas_tabela = ["PTM", 'VALOR TOTAL FEM', "VALOR REPASSADO", "DATA ÚLTIMO PAGAMENTO", 'STATUS']
            
            tabela = doc.add_table(rows=1, cols=len(colunas_tabela))
            tabela.style = 'Table Grid'
            
            # Formata o cabeçalho da tabela (negrito, cor branca e fundo colorido)
            hdr_cells = tabela.rows[0].cells
            for i, coluna in enumerate(colunas_tabela):
                formatar_cabecalho_tabela(hdr_cells[i], coluna)
            
            # Adiciona as linhas da tabela
            for _, linha in df_ano.iterrows():
                row_cells = tabela.add_row().cells
                for i, coluna in enumerate(colunas_tabela):
                    if coluna == "PTM":
                        valor = linha.get("PROJETO DETALHADO", '0')
                    elif coluna == "VALOR TOTAL FEM":
                        valor = linha.get('TETO FEM_x' if 'TETO FEM_x' in df_ano.columns else 'TETO FEM_y', '0')
                        valor = formatar_moeda(valor)
                    elif coluna == "VALOR REPASSADO":
                        valor = linha.get("REPASSE_VÁLIDO", '0')
                        valor = formatar_moeda(valor)
                    elif coluna == "DATA ÚLTIMO PAGAMENTO":
                        valor = str(linha.get("DATA ÚLTIMO PAGAMENTO", '0'))
                        try:
                            valor = datetime.strptime(valor, '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')
                        except ValueError:
                            pass
                    elif coluna == "STATUS":
                        valor = linha.get('STATUS OBRA_x' if 'STATUS OBRA_x' in df_ano.columns else 'STATUS OBRA_y', '0')
                    row_cells[i].text = str(valor)
            
            # Adiciona a linha de total
            total_row = tabela.add_row().cells
            formatar_celula_negrito(total_row[0], "VALOR TOTAL")  # Formata "VALOR TOTAL" em negrito
            
            # Calcula e formata os totais para VALOR TOTAL FEM e VALOR REPASSADO
            coluna_teto_fem = 'TETO FEM_x' if 'TETO FEM_x' in df_ano.columns else 'TETO FEM_y'
            total_teto_fem = df_ano[coluna_teto_fem].replace('', '0').astype(float).sum()
            formatar_celula_negrito(total_row[1], formatar_moeda(total_teto_fem))  # Formata o valor em negrito
            
            total_repasse_valido = df_ano['REPASSE_VÁLIDO'].replace('', '0').astype(float).sum()
            formatar_celula_negrito(total_row[2], formatar_moeda(total_repasse_valido))  # Formata o valor em negrito
            
            # Aplica a formatação da tabela (fonte Arial tamanho 8)
            formatar_tabela(tabela)
        
        # Adiciona a tabela de Emendas Parlamentares
        if not df_emendas_filtrado.empty:
            adicionar_paragrafo_formatado(doc, "Emendas Parlamentares", fonte='Arial', tamanho=12, negrito=True, cor=None)
            
            colunas_tabela_emendas = ["PROJETO DETALHADO", "STATUS OBRA", "VALOR UTILIZADO DA EMENDA", "REPASSE_VÁLIDO", "DATA ÚLTIMO PAGAMENTO"]
            
            tabela_emendas = doc.add_table(rows=1, cols=len(colunas_tabela_emendas))
            tabela_emendas.style = 'Table Grid'
            
            # Formata o cabeçalho da tabela (negrito, cor branca e fundo colorido)
            hdr_cells = tabela_emendas.rows[0].cells
            for i, coluna in enumerate(colunas_tabela_emendas):
                formatar_cabecalho_tabela(hdr_cells[i], coluna)
            
            # Adiciona as linhas da tabela
            for _, linha in df_emendas_filtrado.iterrows():
                row_cells = tabela_emendas.add_row().cells
                for i, coluna in enumerate(colunas_tabela_emendas):
                    if coluna == "VALOR UTILIZADO DA EMENDA" or coluna == "REPASSE_VÁLIDO":
                        valor = formatar_moeda(linha.get(coluna, '0'))
                    else:
                        valor = str(linha.get(coluna, '0'))
                    row_cells[i].text = valor
            
            # Aplica a formatação da tabela (fonte Arial tamanho 8)
            formatar_tabela(tabela_emendas)
        
        doc.save(caminho_nota)
        st.success("✅ Nota Técnica gerada com sucesso!")
        with open(caminho_nota, "rb") as file:
            st.download_button("📥 Baixar Nota Técnica (DOCX)", file, file_name="nota_tecnica.docx")
